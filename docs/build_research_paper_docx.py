from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "research_paper.md"
OUTPUT = ROOT / "docs" / "Customer_Segmentation_Churn_Research_Paper.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(107, 114, 128)
LIGHT_FILL = "F2F4F7"


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    setup_document(doc)
    add_front_matter(doc)
    render_markdown_body(doc, markdown)
    add_footer_page_numbers(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.bold = True
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = BLUE
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)

    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = BLUE
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)

    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = DARK_BLUE
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)


def add_front_matter(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Customer Segmentation and Churn Pattern Analytics in European Banking")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Research Paper Submission")
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    metadata = [
        ("Author", "Shridhar Kalasgonda"),
        ("Domain", "Banking Analytics and Business Intelligence"),
        ("Dataset", "European banking customer churn dataset, 10,000 records"),
        ("Tools", "Python, Pandas, Streamlit, Plotly"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [2200, 6200])
    for label, value in metadata:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], LIGHT_FILL)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].bold = True
    set_table_borders(table)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Prepared for academic and portfolio submission")
    run.font.italic = True
    run.font.color.rgb = MUTED

    doc.add_section(WD_SECTION.NEW_PAGE)


def render_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    # Skip title and author block because the DOCX has a designed cover page.
    while index < len(lines) and not lines[index].startswith("## Abstract"):
        index += 1

    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
        elif line.startswith("# "):
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
        elif re.match(r"^\d+\.\s+", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_runs(paragraph, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, line[2:].strip())
        else:
            paragraph = doc.add_paragraph()
            add_inline_runs(paragraph, line)
        index += 1


def add_inline_runs(paragraph, text: str) -> None:
    text = text.replace("  ", " ")
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = choose_widths(col_count)
    set_table_widths(table, widths)
    set_table_borders(table)

    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            value = row[col_index] if col_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_number_like(value) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(9)
            if row_index == 0:
                shade_cell(cell, LIGHT_FILL)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = INK

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def choose_widths(col_count: int) -> list[int]:
    if col_count == 2:
        return [2600, 6760]
    if col_count == 3:
        return [3000, 2100, 4260]
    if col_count == 4:
        return [2600, 1900, 1900, 2960]
    if col_count == 5:
        return [2100, 1400, 1300, 1500, 3060]
    if col_count == 6:
        return [1800, 1250, 1100, 1250, 1900, 2060]
    return [int(9360 / col_count)] * col_count


def is_number_like(value: str) -> bool:
    cleaned = value.replace("USD", "").replace("$", "").replace(",", "").replace("%", "").strip()
    try:
        float(cleaned)
        return True
    except ValueError:
        return False


def set_table_widths(table, widths: list[int]) -> None:
    table_width = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Pt(width / 20)
                tc_pr = row.cells[index]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(width))
                tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D1D5DB")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def add_footer_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.paragraph_format.space_after = Pt(0)
        footer.add_run("Page ")
        add_field(footer, "PAGE")


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, text, fld_char_end])


if __name__ == "__main__":
    main()
